"""
Document processing service for SCONIA.
Handles text extraction, preprocessing, and document management for legal documents.
"""
from typing import List, Dict, Any, Optional, BinaryIO
import logging
import os
import re
from pathlib import Path
import asyncio
from datetime import datetime
import hashlib

# Document processing libraries
import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
import aiofiles

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update

from app.models.embeddings import ProcessedDocument
from app.services.embeddings import embedding_service
from app.services.vector_db import vector_db_service
from app.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing and extracting text from legal documents."""
    
    def __init__(self):
        """Initialize document processor."""
        self.supported_formats = ['.pdf', '.docx', '.txt', '.html', '.htm']
        self.upload_dir = Path(settings.upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        try:
            # Remove excessive whitespace
            text = re.sub(r'\s+', ' ', text)
            
            # Remove special characters but keep legal punctuation
            text = re.sub(r'[^\w\s\.\,\;\:\!\?\(\)\[\]\-\'\"]', ' ', text)
            
            # Fix common OCR errors in legal documents
            text = re.sub(r'\b(\d+)\s*\.\s*(\d+)\b', r'\1.\2', text)  # Fix section numbers
            text = re.sub(r'\b(Section|Article|Chapter)\s+(\d+)', r'\1 \2', text, flags=re.IGNORECASE)
            
            # Normalize quotes
            text = re.sub(r'["""]', '"', text)
            text = re.sub(r"[''']", "'", text)
            
            # Remove extra spaces
            text = re.sub(r'\s+', ' ', text).strip()
            
            return text
            
        except Exception as e:
            logger.warning(f"Error cleaning text: {e}")
            return text
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1}: {e}")
                        continue
            
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return self._clean_text(text)
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    return self._clean_text(text)
                except UnicodeDecodeError:
                    continue
            
            logger.error(f"Could not decode text file {file_path}")
            raise
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {e}")
            raise
    
    def extract_text_from_html(self, file_path: str) -> str:
        """Extract text from HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting text from HTML {file_path}: {e}")
            raise
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from file based on extension."""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif file_ext == '.docx':
                return self.extract_text_from_docx(file_path)
            elif file_ext == '.txt':
                return self.extract_text_from_txt(file_path)
            elif file_ext in ['.html', '.htm']:
                return self.extract_text_from_html(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    def _generate_document_id(self, file_path: str, content: str) -> str:
        """Generate unique document ID based on file path and content."""
        content_hash = hashlib.md5(content.encode()).hexdigest()
        file_name = Path(file_path).stem
        return f"{file_name}_{content_hash[:8]}"
    
    def _extract_metadata(self, text: str, document_type: str) -> Dict[str, Any]:
        """Extract metadata from document text."""
        metadata = {
            "word_count": len(text.split()),
            "char_count": len(text),
            "document_type": document_type
        }
        
        # Extract legal-specific metadata
        if document_type == "constitution":
            # Extract chapter and section information
            chapters = re.findall(r'Chapter\s+([IVX]+|\d+)', text, re.IGNORECASE)
            sections = re.findall(r'Section\s+(\d+)', text, re.IGNORECASE)
            metadata.update({
                "chapters": list(set(chapters)),
                "sections": list(set(sections))
            })
        
        elif document_type == "case":
            # Extract case information
            case_numbers = re.findall(r'(\d{4})\s*[A-Z]+\s*(\d+)', text)
            judges = re.findall(r'Justice\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)', text)
            metadata.update({
                "case_numbers": case_numbers,
                "judges_mentioned": list(set(judges))
            })
        
        return metadata
    
    async def save_uploaded_file(self, file: BinaryIO, filename: str) -> str:
        """Save uploaded file to disk."""
        try:
            # Generate unique filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_filename = re.sub(r'[^\w\-_\.]', '_', filename)
            unique_filename = f"{timestamp}_{safe_filename}"
            
            file_path = self.upload_dir / unique_filename
            
            # Save file
            async with aiofiles.open(file_path, 'wb') as f:
                content = await file.read()
                await f.write(content)
            
            logger.info(f"Saved uploaded file: {file_path}")
            return str(file_path)
            
        except Exception as e:
            logger.error(f"Error saving uploaded file: {e}")
            raise
    
    async def process_document(
        self,
        file_path: str,
        document_type: str,
        db: AsyncSession,
        use_gemini: bool = True
    ) -> str:
        """
        Process a document: extract text, generate embeddings, and store in vector DB.
        
        Args:
            file_path: Path to the document file
            document_type: Type of document (constitution, case, etc.)
            db: Database session
            use_gemini: Whether to use Gemini for embeddings
            
        Returns:
            Document ID
        """
        try:
            # Extract text
            logger.info(f"Processing document: {file_path}")
            text = self.extract_text(file_path)
            
            if not text.strip():
                raise ValueError("No text extracted from document")
            
            # Generate document ID
            document_id = self._generate_document_id(file_path, text)
            
            # Extract metadata
            metadata = self._extract_metadata(text, document_type)
            metadata.update({
                "file_path": file_path,
                "file_size": os.path.getsize(file_path),
                "processed_at": datetime.utcnow().isoformat()
            })
            
            # Create processed document record
            processed_doc = ProcessedDocument(
                document_id=document_id,
                document_name=Path(file_path).name,
                document_type=document_type,
                file_path=file_path,
                file_size=os.path.getsize(file_path),
                processing_status="processing"
            )
            
            db.add(processed_doc)
            await db.commit()
            
            try:
                # Generate embeddings
                embedding_records = await embedding_service.embed_document(
                    text=text,
                    document_id=document_id,
                    document_type=document_type,
                    metadata=metadata,
                    use_gemini=use_gemini
                )
                
                # Store embeddings in vector database
                embeddings = [record["embedding"] for record in embedding_records]
                texts = [record["text"] for record in embedding_records]
                metadata_list = [record["metadata"] for record in embedding_records]
                
                await vector_db_service.store_embeddings(
                    embeddings=embeddings,
                    texts=texts,
                    metadata_list=metadata_list
                )
                
                # Update processed document status
                await db.execute(
                    update(ProcessedDocument)
                    .where(ProcessedDocument.document_id == document_id)
                    .values(
                        processing_status="completed",
                        chunk_count=len(embedding_records),
                        processed_at=datetime.utcnow()
                    )
                )
                await db.commit()
                
                logger.info(f"Successfully processed document {document_id}")
                return document_id
                
            except Exception as e:
                # Update status to failed
                await db.execute(
                    update(ProcessedDocument)
                    .where(ProcessedDocument.document_id == document_id)
                    .values(
                        processing_status="failed",
                        error_message=str(e)
                    )
                )
                await db.commit()
                raise
                
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            raise
    
    async def get_processed_documents(self, db: AsyncSession) -> List[Dict[str, Any]]:
        """Get list of processed documents."""
        try:
            result = await db.execute(select(ProcessedDocument))
            documents = result.scalars().all()
            
            return [
                {
                    "document_id": doc.document_id,
                    "document_name": doc.document_name,
                    "document_type": doc.document_type,
                    "processing_status": doc.processing_status,
                    "chunk_count": doc.chunk_count,
                    "file_size": doc.file_size,
                    "processed_at": doc.processed_at,
                    "error_message": doc.error_message
                }
                for doc in documents
            ]
            
        except Exception as e:
            logger.error(f"Error getting processed documents: {e}")
            return []


# Global document processor instance
document_processor = DocumentProcessor()
